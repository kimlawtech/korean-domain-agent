"""에이전트 결과(텍스트) → Word(.docx).

python-docx 가 설치돼 있으면 docx 로,
없으면 .md 로 fallback 저장합니다 — 워크숍 환경 호환성.
"""
from __future__ import annotations

from pathlib import Path


def save(text: str, out_path: str | Path) -> str:
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt
    except Exception:
        # python-docx 미설치 → .md 로 저장
        md_path = out.with_suffix(".md")
        md_path.write_text(text, encoding="utf-8")
        return str(md_path)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Pretendard"
    style.font.size = Pt(11)

    for raw in text.split("\n"):
        line = raw.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        else:
            doc.add_paragraph(line)

    doc.save(str(out))
    return str(out)
